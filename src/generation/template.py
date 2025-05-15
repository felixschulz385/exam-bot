from docx import Document
from pathlib import Path
from typing import Dict, Any, Optional, List
import logging
import re

logger = logging.getLogger(__name__)

class TemplateProcessor:
    """Process templates with placeholders for exam generation."""
    
    # Define placeholder patterns
    PLACEHOLDER_PATTERN = r'\{\{\s*([A-Za-z_][A-Za-z0-9_]*)\s*\}\}'
    
    def __init__(self, template_path: Path):
        """Initialize template processor.
        
        Args:
            template_path: Path to Word document template
        """
        self.template_path = Path(template_path)
        
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")
    
    def load_template(self) -> Document:
        """Load the template document.
        
        Returns:
            Document object
        """
        try:
            return Document(self.template_path)
        except Exception as e:
            logger.error(f"Error loading template: {e}")
            raise
    
    def find_placeholders(self, doc: Document) -> List[str]:
        """Find all placeholders in the document.
        
        Args:
            doc: Document object
        
        Returns:
            List of placeholder names
        """
        placeholders = set()
        
        # Search for placeholders in paragraphs
        for para in doc.paragraphs:
            for match in re.finditer(self.PLACEHOLDER_PATTERN, para.text):
                placeholders.add(match.group(1))
        
        # Search for placeholders in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for match in re.finditer(self.PLACEHOLDER_PATTERN, para.text):
                            placeholders.add(match.group(1))
        
        return list(placeholders)
    
    def replace_placeholders(self, 
                            doc: Document, 
                            replacements: Dict[str, Any]) -> Document:
        """Replace all placeholders with provided values.
        
        Args:
            doc: Document object
            replacements: Dictionary mapping placeholder names to replacement values
            
        Returns:
            Modified document
        """
        # Replace placeholders in paragraphs
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)
        
        # Replace placeholders in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)
        
        return doc
    
    def _replace_in_paragraph(self, 
                             para, 
                             replacements: Dict[str, Any]) -> None:
        """Replace placeholders in a single paragraph.
        
        Args:
            para: Paragraph object
            replacements: Dictionary mapping placeholder names to replacement values
        """
        if not para.text:
            return
            
        # Find all placeholders in this paragraph
        matches = list(re.finditer(self.PLACEHOLDER_PATTERN, para.text))
        
        # If no placeholders, do nothing
        if not matches:
            return
            
        # Replace all placeholders from last to first to preserve indices
        for match in reversed(matches):
            placeholder = match.group(1)
            start, end = match.span()
            
            if placeholder in replacements:
                # Get replacement value, convert to string if needed
                replacement = str(replacements[placeholder])
                
                # Replace text in paragraph
                para.text = para.text[:start] + replacement + para.text[end:]
                
            else:
                logger.warning(f"No replacement provided for placeholder: {placeholder}")
    
    def process_template(self, 
                        replacements: Dict[str, Any], 
                        output_path: Optional[Path] = None) -> Path:
        """Process template with replacements and save to output file.
        
        Args:
            replacements: Dictionary mapping placeholder names to replacement values
            output_path: Path to save output document
            
        Returns:
            Path to saved document
        """
        # Load template
        doc = self.load_template()
        
        # Find placeholders and check for missing replacements
        placeholders = self.find_placeholders(doc)
        missing = [p for p in placeholders if p not in replacements]
        
        if missing:
            logger.warning(f"Missing replacements for placeholders: {', '.join(missing)}")
        
        # Replace placeholders
        doc = self.replace_placeholders(doc, replacements)
        
        # Determine output path if not provided
        if not output_path:
            stem = self.template_path.stem
            suffix = self.template_path.suffix
            output_path = self.template_path.parent / f"{stem}_processed{suffix}"
        
        # Save document
        doc.save(output_path)
        logger.info(f"Processed template saved to: {output_path}")
        
        return output_path