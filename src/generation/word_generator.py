from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple
import logging
import re
import datetime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random

logger = logging.getLogger(__name__)

class WordGenerator:
    """Generates Word document exams from selected questions."""
    
    def __init__(self, 
                template_path: Optional[Path] = None,
                output_dir: Path = Path("output")):
        """Initialize Word document generator.
        
        Args:
            template_path: Path to Word document template
            output_dir: Directory for output files
        """
        self.template_path = template_path
        self.output_dir = Path(output_dir)
        self.doc_config = {}  # Store document config
        
        # Create output directory if it doesn't exist
        if not self.output_dir.exists():
            self.output_dir.mkdir(parents=True)
            logger.info(f"Created output directory: {self.output_dir}")
    
    def _create_document(self) -> Document:
        """Create a new Word document from template or blank.
        
        Returns:
            Document object
        """
        if self.template_path and Path(self.template_path).exists():
            try:
                doc = Document(self.template_path)
                logger.info(f"Created document from template: {self.template_path}")
                return doc
            except Exception as e:
                logger.error(f"Error loading template: {e}")
                logger.info("Falling back to blank document")
        
        # Fall back to blank document
        return Document()
    
    def _apply_document_styles(self, 
                              doc: Document, 
                              config: Dict[str, Any]) -> None:
        """Apply document styles based on configuration.
        
        Args:
            doc: Document object
            config: Configuration dictionary with formatting options
        """
        format_config = config.get("format", {})
        
        # Set default font and size for the document
        font_name = format_config.get("font", "Calibri")
        font_size = format_config.get("font_size", 11)
        
        styles = doc.styles
        
        # Modify the Normal style
        style_normal = styles['Normal']
        font = style_normal.font
        font.name = font_name
        font.size = Pt(font_size)
        
        # Create or modify heading styles
        for level in range(1, 4):
            style_name = f'Heading {level}'
            if style_name in styles:
                style = styles[style_name]
            else:
                style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                
            style.base_style = styles['Normal']
            style.font.name = font_name
            style.font.bold = True
            style.font.size = Pt(font_size + (4 - level))
        
        # Create a style for question text
        if 'Question' not in styles:
            question_style = styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
            question_style.base_style = styles['Normal']
            question_style.font.bold = True
        else:
            question_style = styles['Question']
        
        # Create a style for answer choices
        if 'Answer' not in styles:
            answer_style = styles.add_style('Answer', WD_STYLE_TYPE.PARAGRAPH)
            answer_style.base_style = styles['Normal']
            answer_style.paragraph_format.left_indent = Inches(0.25)
        
        # Create a style for correct answers (instructor version)
        if 'CorrectAnswer' not in styles:
            correct_style = styles.add_style('CorrectAnswer', WD_STYLE_TYPE.CHARACTER)
            correct_style.font.color.rgb = RGBColor(0, 128, 0)  # Green
            correct_style.font.bold = True
    
        # Create a style for block questions
        if 'BlockQuestion' not in styles:
            block_style = styles.add_style('BlockQuestion', WD_STYLE_TYPE.PARAGRAPH)
            block_style.base_style = styles['Question']
            block_style.font.italic = True
            # Add light gray background for block questions
            if hasattr(block_style.paragraph_format, 'shading'):
                block_style.paragraph_format.shading.fill = 'F0F0F0'  # Light gray
    
        # Create a style for block headings
        if 'BlockHeading' not in styles:
            block_heading_style = styles.add_style('BlockHeading', WD_STYLE_TYPE.PARAGRAPH)
            block_heading_style.base_style = styles['Heading 3']
            block_heading_style.font.bold = True
            block_heading_style.font.color.rgb = RGBColor(50, 50, 150)  # Navy blue
    
    def _add_header_footer(self, 
                          doc: Document, 
                          config: Dict[str, Any]) -> None:
        """Add header and footer to document.
        
        Args:
            doc: Document object
            config: Configuration dictionary with header/footer text
        """
        header_text = config.get("header", "")
        footer_text = config.get("footer", "")
        
        if header_text:
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = header_text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if footer_text:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_title_and_instructions(self, 
                                  doc: Document, 
                                  config: Dict[str, Any]) -> None:
        """Add title and instructions to document.
        
        Args:
            doc: Document object
            config: Configuration dictionary with title and instructions
        """
        title_text = config.get("title", "Exam")
        instructions = config.get("instructions", "")
        
        # Add title
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
        
        # Add instructions if provided
        if instructions:
            doc.add_paragraph()  # Add blank line
            instructions_para = doc.add_paragraph(instructions)
            doc.add_paragraph()  # Add blank line after instructions
    
    def _format_question_text(self, text: str) -> str:
        """Format question text by removing excess whitespace and normalizing.
        
        Args:
            text: Raw question text
            
        Returns:
            Formatted question text
        """
        # Remove excess whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def _generate_version(self, 
                    questions: pd.DataFrame, 
                    config: Dict[str, Any],
                    include_answers: bool = True,
                    mark_correct: bool = False,
                    version_suffix: str = "") -> Tuple[Path, Dict]:
        """Generate a specific version of the exam.
        
        Args:
            questions: DataFrame of selected questions
            config: Configuration dictionary
            include_answers: Whether to include answer choices
            mark_correct: Whether to mark correct answers
            version_suffix: Suffix to add to filename
            
        Returns:
            Tuple of (path to generated document, answer mapping dictionary)
        """
        # Store document config for use in other methods
        self.doc_config = config.get("document", {})
        
        # Check if we should randomize answer options
        randomize_answers = self.doc_config.get("randomize_answers", False)
        self.seed = config.get("selection", {}).get("seed")
        
        # Store answer mappings for each question (for marking sheet)
        answer_mappings = {}
        
        # Create new document
        doc = self._create_document()
        
        # Apply styles
        self._apply_document_styles(doc, self.doc_config)
        
        # Add page numbers
        self._add_page_numbers(doc)
        
        # Add custom header if provided
        header_text = config.get("exam", {}).get("header_text", "")
        if header_text:
            self._add_custom_header(doc, header_text)
            
        # Add exam instructions if provided
        instructions = config.get("exam", {}).get("instructions", "")
        if instructions:
            self._add_instructions(doc, instructions)
        
        # Get question bank reader for block information
        question_bank_reader = config.get("_runtime", {}).get("question_bank_reader")
        has_block_info = question_bank_reader is not None
        logger.info(f"Has question bank reader for block info: {has_block_info}")
        
        # Get type points for adding to questions
        grading_config = config.get("grading", {})
        type_points = grading_config.get("type_points", {})
        
        # Skip topic and type headers as requested
        include_topic_headers = False  # Explicitly set to false
        include_type_headers = False   # Explicitly set to false
        
        # Add questions to document
        question_num = 1  # Keep this as 1 for internal logic
        current_block = None
        figure_num = 1
        total_points = 0
        
        # First pass: identify last question in each block
        last_in_block = {}
        if has_block_info:
            block_questions = {}
            # Group questions by block
            for idx, row in questions.iterrows():
                try:
                    is_block_question = question_bank_reader.is_block_question.get(idx, False)
                    if is_block_question:
                        block_id = question_bank_reader.get_block_for_question(idx)
                        if block_id not in block_questions:
                            block_questions[block_id] = []
                        block_questions[block_id].append(idx)
                except Exception:
                    pass
                    
            # Mark last question in each block
            for block_id, question_indices in block_questions.items():
                if question_indices:
                    last_in_block[question_indices[-1]] = True

        # Process questions
        for idx, row in questions.iterrows():
            # Check if this is a block question
            is_block_question = False
            is_first_in_block = False
            is_last_in_block = last_in_block.get(idx, False)

            if has_block_info:
                try:
                    # Try accessing is_block_question as a dictionary
                    is_block_question = question_bank_reader.is_block_question.get(idx, False)
                
                    if is_block_question:
                        block_id = question_bank_reader.get_block_for_question(idx)
                        is_first_in_block = question_bank_reader.is_first_in_block(idx)
                
                        # If this is the first question of a new block, add strong visual separation
                        if block_id != current_block and is_first_in_block:
                            current_block = block_id
                            if question_num > 1:  # Don't add extra space at the beginning
                                # Add stronger block separation
                                self._add_block_separator(doc)
                                
                            # Add block heading with ID
                            self._format_block_heading(doc, block_id)
                except Exception as e:
                    logger.warning(f"Block question detection failed: {e}, treating as non-block")
        
            # Calculate points for this question
            points = type_points.get(row['Type'], 1)
            total_points += points

            # Add the question with points information
            block_info = {
                'is_block_question': is_block_question,
                'is_first_in_block': is_first_in_block,
                'is_last_in_block': is_last_in_block,
                'points': points,
                'figure_num': figure_num
            }
            
            # If this question has a figure, increment the figure number for next time
            if 'Medium' in row and pd.notna(row['Medium']):
                if re.search(r'(?:figure|fig\.?)\s+(\w+)', row['Medium'], re.IGNORECASE):
                    figure_num += 1
            
            # Add the question (the internal logic stays at 1-based,
            # but _add_question will display it as 0-based with 2 digits)
            answer_mapping = self._add_question(
                doc, question_num, row, include_answers, mark_correct, block_info,
                randomize_answers=randomize_answers
            )
            
            # Store the mapping for this question
            if randomize_answers:
                answer_mappings[question_num] = answer_mapping
            
            question_num += 1

        # Add total points summary at the end
        self._add_points_summary(doc, total_points)

        # Set up path for output file
        exam_title = config.get("exam", {}).get("title", "Exam").replace(" ", "_")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        if version_suffix:
            filename = f"{exam_title}_{version_suffix}_{timestamp}.docx"
        else:
            filename = f"{exam_title}_{timestamp}.docx"
            
        output_path = self.output_dir / filename

        # Save the document
        doc.save(output_path)
        logger.info(f"Generated exam document: {output_path}")
        
        return output_path, answer_mappings

    def _add_custom_header(self, doc: Document, header_text: str) -> None:
        """Add a custom header to the document.
    
        Args:
            doc: Document object
            header_text: Header text to display
        """
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run(header_text)
        header_run.bold = True
        header_run.font.size = Pt(14)
        
        # Add a separator line below the header
        doc.add_paragraph("_" * 80)
        
        # Add some space after header
        doc.add_paragraph()
        
        logger.info(f"Added custom header: {header_text}")

    def _add_points_summary(self, doc: Document, total_points: int) -> None:
        """Add total points summary at the end of the document.
    
        Args:
            doc: Document object
            total_points: Total points possible in the exam
        """
        # Add some separation before summary
        doc.add_paragraph("_" * 80)
    
        # Add total points information
        points_para = doc.add_paragraph()
        points_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        points_run = points_para.add_run(f"Total possible points: {total_points}")
        points_run.bold = True
        points_run.font.size = Pt(12)
    
        logger.info(f"Added points summary: {total_points} total points")

    def _add_page_numbers(self, doc: Document) -> None:
        """Add simple running page numbers to the document.
        
        Args:
            doc: Document object
        """
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        
        # Clear any existing content
        footer_para.clear()
        
        # Center align the paragraph
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            # Try to use the XML approach to add page numbers
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # Add "Page " text
            footer_para.add_run("Page ")
            
            # Add page field
            run = footer_para.add_run()
            
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = " PAGE "
            run._r.append(instrText)
            
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)
            
            logger.info("Added page numbers using field codes")
            
        except (ImportError, AttributeError) as e:
            # Fallback to placeholder if the XML approach doesn't work
            footer_para.text = "Page [Page #]"
            logger.warning(f"Could not add automatic page numbers: {e}. Using placeholder instead.")

    def _add_front_matter_document(self, doc: Document, front_matter_path: Path) -> None:
        """Add content from a front matter document to the beginning of an exam.
    
        Args:
            doc: Target Document object
            front_matter_path: Path to Word document with front matter
        """
        if not front_matter_path.exists():
            logger.warning(f"Front matter document not found: {front_matter_path}")
            return
            
        try:
            # Open the front matter document
            front_matter_doc = Document(front_matter_path)
            
            # Get all paragraphs from the front matter
            paragraphs = list(front_matter_doc.paragraphs)
            
            # Get all tables from the front matter
            tables = list(front_matter_doc.tables)
            
            # Create temporary document to hold the exam content
            temp_doc = Document()
            
            # Copy all styles from the original document
            # (We need to do this to ensure formatting is consistent)
            for style in doc.styles:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    if style.name not in temp_doc.styles:
                        try:
                            temp_doc.styles.add_style(
                                style.name, 
                                WD_STYLE_TYPE.PARAGRAPH
                            ).base_style = temp_doc.styles['Normal']
                        except:
                            # If style already exists or can't be created, skip
                            pass
            
            # Copy all content from the original document to the temp document
            for p in doc.paragraphs:
                new_p = temp_doc.add_paragraph()
                new_p.text = p.text
                new_p.style = p.style.name
                
                # Copy runs with their formatting
                for run in p.runs:
                    new_run = new_p.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    
            # Now clear the original document
            for p in list(doc.paragraphs):
                p.clear()
                
            # Add front matter
            # First, copy paragraphs from front matter
            for p in paragraphs:
                new_p = doc.add_paragraph()
                new_p.text = p.text
                
                # Copy runs with their formatting
                for run in p.runs:
                    new_run = new_p.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    
            # Copy tables from front matter if any
            for table in tables:
                # Get row and column count
                rows = len(table.rows)
                cols = len(table.rows[0].cells) if rows > 0 else 0
                
                if rows > 0 and cols > 0:
                    # Create a new table with the same dimensions
                    new_table = doc.add_table(rows=rows, cols=cols)
                    
                    # Copy content from each cell
                    for i in range(rows):
                        for j in range(cols):
                            src_cell = table.cell(i, j)
                            tgt_cell = new_table.cell(i, j)
                            tgt_cell.text = src_cell.text
                
                # Add space after table
                doc.add_paragraph()
                
            # Add a page break between front matter and exam content
            doc.add_page_break()
            
            # Copy back all content from the temp document
            for p in temp_doc.paragraphs:
                new_p = doc.add_paragraph()
                new_p.text = p.text
                
                # Copy runs with their formatting
                for run in p.runs:
                    new_run = new_p.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    
            logger.info(f"Added front matter from: {front_matter_path}")
            
        except Exception as e:
            logger.error(f"Error adding front matter: {e}")
            # Continue without front matter rather than failing

    def convert_to_pdf(self, docx_path: Path, delete_original: bool = False) -> Optional[Path]:
        """Convert a Word document to PDF.
    
        Args:
            docx_path: Path to Word document
            delete_original: Whether to delete the original docx file
            
        Returns:
            Path to generated PDF, or None if conversion failed
        """
        if not docx_path.exists():
            logger.warning(f"Document not found for PDF conversion: {docx_path}")
            return None
        
        pdf_path = docx_path.with_suffix(".pdf")
        
        try:
            # First try using LibreOffice (works on many platforms)
            try:
                # Check if libreoffice or soffice is available
                if os.name == 'nt':  # Windows
                    libreoffice_commands = ['soffice.exe']
                else:  # macOS, Linux
                    libreoffice_commands = ['libreoffice', 'soffice']
                    
                command_exists = False
                command_to_use = None
                
                for cmd in libreoffice_commands:
                    try:
                        subprocess.run(['which', cmd], check=True, capture_output=True)
                        command_exists = True
                        command_to_use = cmd
                        break
                    except:
                        pass
                        
                if command_exists:
                    # Convert using LibreOffice
                    subprocess.run([
                        command_to_use,
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', str(docx_path.parent),
                        str(docx_path)
                    ], check=True, capture_output=True)
                    
                    logger.info(f"Converted to PDF using LibreOffice: {pdf_path}")
                    
                    if delete_original:
                        docx_path.unlink()
                        logger.info(f"Deleted original docx: {docx_path}")
                        
                    return pdf_path
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed, trying alternative: {e}")
            
            # If LibreOffice failed, try using docx2pdf if available
            try:
                from docx2pdf import convert
                convert(str(docx_path), str(pdf_path))
                logger.info(f"Converted to PDF using docx2pdf: {pdf_path}")
                
                if delete_original:
                    docx_path.unlink()
                    logger.info(f"Deleted original docx: {docx_path}")
                    
                return pdf_path
            except ImportError:
                logger.warning("docx2pdf not installed, trying another method")
            
            # If all else fails, notify the user
            logger.warning("PDF conversion failed. Please install LibreOffice or docx2pdf package.")
            logger.warning("You can install docx2pdf with: pip install docx2pdf")
            return None
            
        except Exception as e:
            logger.error(f"Error during PDF conversion: {e}")
            return None

    def combine_pdfs(self, exam_pdf_path: Path, front_matter_pdf_path: Path) -> Optional[Path]:
        """Combine front matter PDF with exam PDF.
    
        Args:
            exam_pdf_path: Path to exam PDF file
            front_matter_pdf_path: Path to front matter PDF file
            
        Returns:
            Path to combined PDF, or None if combination failed
        """
        try:
            # Import PyPDF2 for PDF operations
            try:
                from PyPDF2 import PdfMerger
            except ImportError:
                logger.error("PyPDF2 package not found. Install with: pip install PyPDF2")
                return None
                
            if not exam_pdf_path.exists():
                logger.error(f"Exam PDF not found: {exam_pdf_path}")
                return None
                
            if not front_matter_pdf_path.exists():
                logger.error(f"Front matter PDF not found: {front_matter_pdf_path}")
                return None
                
            # Create output path - same name but with _complete suffix
            output_path = exam_pdf_path.parent / (exam_pdf_path.stem + "_complete.pdf")
            
            # Create PDF merger
            merger = PdfMerger()
            
            # Add front matter first
            merger.append(str(front_matter_pdf_path))
            
            # Add exam PDF
            merger.append(str(exam_pdf_path))
            
            # Write to output file
            with open(output_path, 'wb') as output_file:
                merger.write(output_file)
                
            logger.info(f"Created combined PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error combining PDFs: {e}")
            return None

    def generate_exam(self, 
                questions: pd.DataFrame, 
                config: Dict[str, Any]) -> Tuple[Optional[Path], Optional[Path], Dict]:
        """Generate exam documents based on selected questions and configuration.
        
        Args:
            questions: DataFrame of selected questions
            config: Configuration dictionary
            
        Returns:
            Tuple of (student_version_path, instructor_version_path, answer_mappings)
        """
        # Don't use the exam title or configuration for file naming - simplified approach
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        
        student_path = None
        instructor_path = None
        answer_mappings = {}
        
        # Generate student version if requested
        if config.get("document", {}).get("student_version", True):
            student_path, answer_mappings = self._generate_version(
                questions, 
                config,
                include_answers=True,
                mark_correct=False,
                version_suffix="student"
            )
        
        # Generate instructor version if requested
        if config.get("document", {}).get("instructor_version", True):
            # Use the same answer mappings for instructor version to keep them consistent
            instructor_path, _ = self._generate_version(
                questions, 
                config,
                include_answers=True,
                mark_correct=True,
                version_suffix="instructor"
            )
        
        return student_path, instructor_path, answer_mappings

    def _add_question(self, 
                doc: Document, 
                question_num: int, 
                question_data: pd.Series,
                include_answers: bool = True,
                mark_correct: bool = False,
                block_info: Optional[Dict] = None,
                randomize_answers: bool = False,
                answer_mapping: Optional[Dict] = None) -> Dict:
        """Add a question to the document."""
        # Handle block question formatting
        is_block_question = False
        is_first_in_block = False
        is_last_in_block = False
        points = 1  # Default to 1 point
        figure_num = block_info.get('figure_num', 1) if block_info else 1
        
        if block_info:
            is_block_question = block_info.get('is_block_question', False)
            is_first_in_block = block_info.get('is_first_in_block', False)
            is_last_in_block = block_info.get('is_last_in_block', False)
            points = block_info.get('points', 1)  # Get points from block_info
    
        # Add block-specific content for the first question in a block
        if is_block_question and is_first_in_block:
            # Add leading text if present - with enhanced formatting
            if 'Leading Text' in question_data and pd.notna(question_data['Leading Text']):
                leading_text = str(question_data['Leading Text'])
                leading_para = doc.add_paragraph()
                leading_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                leading_run = leading_para.add_run(leading_text)
                leading_run.italic = True  # Make leading text italic for visual distinction
                
                # Add horizontal separator after leading text
                doc.add_paragraph("_" * 40)
                logger.info(f"Added leading text for question block with enhanced formatting")
                
            # Handle images without captioning
            if 'Medium' in question_data and pd.notna(question_data['Medium']):
                medium = question_data['Medium']  # e.g., "Figure X"
                
                # Try to extract figure number for file lookup only (not for display)
                try:
                    match = re.search(r'(?:figure|fig\.?)\s+(\w+)', medium, re.IGNORECASE)
                    if match:
                        # Get image directory and original figure number for file lookup
                        original_figure_num = match.group(1)
                        image_dir = self.doc_config.get("block_questions", {}).get("image_directory", "data/Figures")
                        image_path = Path(image_dir) / f"{original_figure_num}.png"
                        
                        # Check if image exists - add image without figure caption
                        if image_path.exists():
                            # Center the image
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(str(image_path), width=Inches(5))
                            logger.info(f"Added figure (from {original_figure_num}.png) without caption")
                        else:
                            logger.warning(f"Figure not found: {image_path}")
                except Exception as e:
                    logger.error(f"Error processing figure: {e}")
    
        # Format question text with points
        question_text = self._format_question_text(str(question_data['Question']))
        
        # For block questions after the first, add indentation and visual distinction
        question_para = doc.add_paragraph(style='Question')
    
        if is_block_question:
            if not is_first_in_block:
                # Increased indentation for follow-up questions in block
                question_para.paragraph_format.left_indent = Inches(0.35)
            else:
                # First question in block gets a subtle highlight
                question_para.paragraph_format.left_indent = Inches(0.2)
    
        # Create two-digit question number format (00, 01, 02, etc.)
        # Subtract 1 from question_num if we want to start from 00 instead of 01
        formatted_question_num = f"{question_num - 1:02d}"
        
        # Add question number and text with points in parentheses    
        question_para.add_run(f"{formatted_question_num}. ").bold = True
        question_para.add_run(question_text)
        question_para.add_run(f" ({points} point{'s' if points > 1 else ''})").italic = True
        
        # Track mapping of original answer positions to randomized positions
        # If not provided, create a new mapping for this question
        if answer_mapping is None:
            answer_mapping = {}
        
        # Add answers if requested
        if include_answers:
            answer_choices = []
            
            # Get answer choices, skipping empty ones
            for i in range(1, 5):
                answer_key = f'Answer {i}'
                if answer_key in question_data and pd.notna(question_data[answer_key]):
                    answer = question_data[answer_key]
                    # Convert answer to string to prevent "int is not iterable" error
                    answer = str(answer)
                    is_correct = False
                    if 'Correct' in question_data:
                        is_correct = i == question_data['Correct']
                    answer_choices.append((i, chr(64 + i), answer, is_correct))
        
            # Randomize answer order if requested
            if randomize_answers:
                # Use a consistent seed based on question_num and an optional seed from config
                # This ensures the randomization is consistent across student/instructor versions
                random_seed = question_num
                if hasattr(self, 'seed') and self.seed is not None:
                    random_seed = random_seed + self.seed
                    
                # Create a random generator for this question
                question_rng = random.Random(random_seed)
                
                # Shuffle the answers
                question_rng.shuffle(answer_choices)
                
                # Create mapping of original positions to new positions
                # Key: original position (1-indexed), Value: new position (letter A, B, C, D)
                for new_idx, (orig_idx, _, _, _) in enumerate(answer_choices):
                    answer_mapping[orig_idx] = chr(65 + new_idx)  # A, B, C, D
        
            # Add each answer choice
            for idx, (orig_idx, letter, answer_text, is_correct) in enumerate(answer_choices):
                answer_para = doc.add_paragraph(style='Answer')
                
                # Add extra indentation for block questions
                if is_block_question:
                    answer_para.paragraph_format.left_indent = Inches(0.5 if is_first_in_block else 0.75)
                
                # Use randomized letter if applicable, otherwise use original letter
                display_letter = chr(65 + idx) if randomize_answers else letter
                
                if mark_correct and is_correct:
                    # For instructor version, mark the correct answer
                    answer_para.add_run(f"{display_letter}) ").bold = True
                    answer_run = answer_para.add_run(answer_text)
                    answer_run.style = 'CorrectAnswer'
                else:
                    answer_para.add_run(f"{display_letter}) {answer_text}")

        # Add appropriate spacing after questions and end of block separator
        if is_block_question:
            if not is_first_in_block and not is_last_in_block:
                # Less space between questions in the same block
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            elif is_last_in_block:
                # Add a simple horizontal separator at the end of a block
                doc.add_paragraph()
                
                # Add a simple horizontal line
                separator_para = doc.add_paragraph()
                separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                separator_para.add_run("─" * 50)
                
                # Add extra space after block
                doc.add_paragraph()
                
                logger.info("Added end-of-block separator")
            else:
                # Normal spacing for the first question
                doc.add_paragraph()
        else:
            # Normal spacing for non-block questions
            doc.add_paragraph()
            
        return answer_mapping

    def _add_block_separator(self, doc: Document) -> None:
        """Add visual separation before a block of questions - now just adds spacing.
    
        Args:
            doc: Document object
        """
        # Just add extra space before block, no horizontal line
        doc.add_paragraph()
        doc.add_paragraph()
        
        logger.info("Added block separator (space only)")

    def _format_block_heading(self, doc: Document, block_id: str) -> None:
        """Process a question block without adding any heading or separator.
    
        Args:
            doc: Document object
            block_id: Identifier for the block (not displayed)
        """
        # No heading or line - just log that we processed the block
        logger.info(f"Processed block {block_id}")

    def _add_instructions(self, doc: Document, instructions: str) -> None:
        """Add exam instructions to the document.
    
        Args:
            doc: Document object
            instructions: Instruction text with markdown formatting
        """
        if not instructions:
            return
            
        # Add title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("EXAM INSTRUCTIONS")
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Add the instructions, handling markdown-style bullet points
        lines = instructions.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                # Empty line - add paragraph break
                doc.add_paragraph()
            elif line.startswith('###'):
                # Section heading
                heading = line.replace('###', '').strip()
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(heading)
                heading_run.bold = True
                heading_run.font.size = Pt(12)
            elif line.startswith('•') or line.startswith('*'):
                # Bullet point
                bullet_text = line[1:].strip()
                bullet_para = doc.add_paragraph(style='ListBullet')
                bullet_para.add_run(bullet_text)
            else:
                # Regular paragraph
                para = doc.add_paragraph()
                para.add_run(line)
    
        # Add separator after instructions
        doc.add_paragraph()
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator.add_run("─" * 40)
        doc.add_paragraph()
        
        logger.info("Added exam instructions")